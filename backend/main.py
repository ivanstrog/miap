import datetime
import os
import shutil
from zipfile import ZipFile

import uvicorn
from fastapi import FastAPI, APIRouter, Depends, status, HTTPException
from typing import Optional

from starlette.middleware.cors import CORSMiddleware

from openpyxl import Workbook

from docx import Document
from docx.shared import Inches

from database_adapter.models import BasePost, PostSeries
from threading import Thread

from database_adapter.post_adapter import PostDatabaseAdapter
from fastapi.responses import FileResponse

from parser_init import update

app = FastAPI(title='Miap Api')

origins = [
    "http://localhost.tiangolo.com",
    "https://localhost.tiangolo.com",
    "http://localhost",
    "http://localhost:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {"message": "Hello World"}


@app.post('/create/')
def create_post(
        post_model: BasePost,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    post_id = post_adapter.create_post(post_model=post_model)

    return {"id": post_id, "status": status.HTTP_201_CREATED}


@app.post('/delete/')
def delete_post(
        post_id: int,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    post_adapter.delete_post(post_id=post_id)
    return status.HTTP_201_CREATED


@app.post('/archive/')
def archive_post(
        post_id: int,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    post_adapter.archive_post(post_id=post_id)
    return status.HTTP_201_CREATED


@app.get('/posts/', response_model=PostSeries)
def get_posts(
        left: int,
        right: int,
        search_company_name: Optional[str] = None,
        search_resource: Optional[str] = None,
        search_title: Optional[str] = None,
        search_category: Optional[str] = None,
        search_link: Optional[str] = None,
        search_doc: Optional[str] = None,
        min_time: Optional[int] = None,
        max_time: Optional[int] = None,
        archive: Optional[bool] = False,
        post_adapter: PostDatabaseAdapter = Depends(),
) -> PostSeries:
    return post_adapter.get_posts_from_l_to_r(left=left, right=right,
                                              search_company_name=search_company_name,
                                              search_resource=search_resource,
                                              search_title=search_title,
                                              search_category=search_category,
                                              search_link=search_link,
                                              search_doc=search_doc,
                                              min_time=min_time,
                                              max_time=max_time,
                                              archive=archive)


@app.get('/get_xlsx_file/')
def get_xlsx(
        left: int,
        right: int,
        search_company_name: Optional[str] = None,
        search_resource: Optional[str] = None,
        search_title: Optional[str] = None,
        search_category: Optional[str] = None,
        search_link: Optional[str] = None,
        search_doc: Optional[str] = None,
        min_time: Optional[int] = None,
        max_time: Optional[int] = None,
        archive: Optional[bool] = False,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    posts_data = post_adapter.get_posts_from_l_to_r(left=left, right=right,
                                                    search_company_name=search_company_name,
                                                    search_resource=search_resource,
                                                    search_title=search_title,
                                                    search_category=search_category,
                                                    search_link=search_link,
                                                    search_doc=search_doc,
                                                    min_time=min_time,
                                                    max_time=max_time,
                                                    archive=archive)

    file_path = 'data.xlsx'

    workbook = Workbook()

    # Get active worksheet/tab
    worksheet = workbook.active
    worksheet.title = 'data'

    # Define the titles for columns
    columns = [
        'ID',
        'Наименование  предприятия  организации',
        'Дата новости  (информации)',
        'Наименование  информационного  ресурса',
        'Наименование  заголовка новости  (информации)',
        'Ссылка на новость  (информацию)',
        'Категория   инвестиционной  активности'
    ]
    row_num = 1

    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value = column_title

    # Iterate through all movies

    for post in posts_data.series:
        row_num += 1
        # Define the data for each cell in the row
        print(type(post))
        row = [
            post.id,
            post.company_name,
            post.date,
            post.resource,
            post.title,
            post.link,
            post.category,
        ]

        # Assign the data for each cell of the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value

    worksheet.column_dimensions['A'].width = 5.0
    worksheet.column_dimensions['B'].width = 40.0
    worksheet.column_dimensions['C'].width = 7.0
    worksheet.column_dimensions['D'].width = 80.0
    worksheet.column_dimensions['E'].width = 80.0
    worksheet.column_dimensions['F'].width = 80.0
    worksheet.column_dimensions['G'].width = 80.0

    worksheet.row_dimensions[1].height = 50

    workbook.save(file_path)
    return FileResponse(media_type='application/octet-stream', filename=file_path, path=file_path)


@app.get('/get_docx_file/')
def get_docx(
        left: int,
        right: int,
        search_company_name: Optional[str] = None,
        search_resource: Optional[str] = None,
        search_title: Optional[str] = None,
        search_category: Optional[str] = None,
        search_link: Optional[str] = None,
        search_doc: Optional[str] = None,
        min_time: Optional[int] = None,
        max_time: Optional[int] = None,
        archive: Optional[bool] = False,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    posts_data = post_adapter.get_posts_from_l_to_r(left=left, right=right,
                                                    search_company_name=search_company_name,
                                                    search_resource=search_resource,
                                                    search_title=search_title,
                                                    search_category=search_category,
                                                    search_link=search_link,
                                                    search_doc=search_doc,
                                                    min_time=min_time,
                                                    max_time=max_time,
                                                    archive=archive)

    document = Document()

    p = document.add_paragraph('Мониторинг инновационной активности предприятий Оренбургской области')

    table = document.add_table(rows=posts_data.number_of_posts + 1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "id"
    hdr_cells[1].text = 'Наименование предприятия организации'
    hdr_cells[2].text = 'Дата новости'
    hdr_cells[3].text = 'Наименование информационного ресурса'
    hdr_cells[4].text = 'Наименование заголовка новости'
    hdr_cells[5].text = 'Ссылка на новость (информацию)'
    hdr_cells[6].text = 'Категория инвестиционной активности'

    for j in range(posts_data.number_of_posts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(posts_data.series[j].id)
        row_cells[1].text = str(posts_data.series[j].company_name)
        row_cells[2].text = str(posts_data.series[j].date)
        row_cells[3].text = str(posts_data.series[j].resource)
        row_cells[4].text = str(posts_data.series[j].title)
        row_cells[5].text = str(posts_data.series[j].link)
        row_cells[6].text = str(posts_data.series[j].category)

    document.add_page_break()

    file_path = 'data.docx'
    document.save(file_path)

    return FileResponse(media_type='application/octet-stream', filename=file_path, path=file_path)


@app.get('/get_zip_file/')
def get_zip(
        left: int,
        right: int,
        search_company_name: Optional[str] = None,
        search_resource: Optional[str] = None,
        search_title: Optional[str] = None,
        search_category: Optional[str] = None,
        search_link: Optional[str] = None,
        search_doc: Optional[str] = None,
        min_time: Optional[int] = None,
        max_time: Optional[int] = None,
        archive: Optional[bool] = False,
        post_adapter: PostDatabaseAdapter = Depends(),
):
    posts_data = post_adapter.get_posts_from_l_to_r(left=left, right=right,
                                                    search_company_name=search_company_name,
                                                    search_resource=search_resource,
                                                    search_title=search_title,
                                                    search_category=search_category,
                                                    search_link=search_link,
                                                    search_doc=search_doc,
                                                    min_time=min_time,
                                                    max_time=max_time,
                                                    archive=archive)

    try:
        os.chdir("backend")
    except Exception as e:
        print(e)
    dir = 'data'
    if os.path.exists(dir):
        shutil.rmtree(dir)
    else:
        print("Текущая деректория:", os.getcwd())
        print('не существует')
    os.makedirs(dir)

    file_path = 'data.zip'
    with ZipFile(file_path, 'w') as myzip:
        for post in posts_data.series:
            filename = f"{dir}/{post.company_name}_{post.category}_{datetime.datetime.now()}.txt"
            my_file = open(filename, "w+")
            my_file.write(f"id в базе данных :    {post.id}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Наименование предприятия организации :   {post.company_name}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Дата новости :  {post.date}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Наименование информационного ресурса :{post.resource}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Наименование заголовка новости :    {post.title}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Ссылка на новость (информацию)    {post.link}")
            my_file.write('\n')
            my_file.write('\n')
            my_file.write(f"Категория инвестиционной активности    {post.category}")
            my_file.write('\n')
            my_file.write('\n')

            try:
                with open(f"storage/{post.link.replace('/','').replace(':','')[-15:]}.html",'r') as html_file:
                    html_code = html_file.read()
                    my_file.write(f"\n Подтверждающие документы( html код страницы): \n    {html_code}")

            except Exception as e:
                print(e)
            my_file.close()
            myzip.write(filename)

    return FileResponse(media_type='application/octet-stream', filename=file_path, path=file_path)


@app.get('/get_post/', response_model=BasePost)
def get_post_by_id(post_id: int, post_adapter: PostDatabaseAdapter = Depends()):
    return post_adapter.get_post_by_id(post_id=post_id)


@app.get('/update_database/')
def get_post_by_id():
    t = Thread(target=update)
    t.start()
    return "обновление запустилось"


if __name__ == '__main__':
    uvicorn.run(
        '__main__:app',
        host='0.0.0.0',
        port=8000,
        reload=True,
    )
